from docx import Document as docxDocument

def convert_docx_to_html(docx_file_path, html_file_path):
    """Converts a .docx file to a simple HTML file."""
    # Load the .docx file
    doc = docxDocument(docx_file_path)

    # Start building the HTML content
    html_content = ["<html>", "<head><title>Converted Document</title></head>", "<body>"]

    # Iterate over the elements in the documents
    for paragraph in doc.paragraphs:
        # Handle headings based on their style
        if paragraph.style.name.startswith("Heading"):
            level = int(paragraph.style.name[-1])  # Extract heading level
            html_content.append(f"<h{level}>{paragraph.text}</h{level}>")
        else:
            # For normal paragraphs
            html_content.append(f"<p>{paragraph.text}</p>")

    # Process tables in the document
    for table in doc.tables:
        html_content.append("<table border='1'>")
        for row in table.rows:
            html_content.append("<tr>")
            for cell in row.cells:
                html_content.append(f"<td>{cell.text}</td>")
            html_content.append("</tr>")
        html_content.append("</table>")

    # Close HTML tags
    html_content.extend(["</body>", "</html>"])

    # Write the HTML content to a file
    with open(html_file_path, "w", encoding="utf-8") as html_file:
        html_file.write("\n".join(html_content))

    # print(f"HTML file created at: {html_file_path}")
    return html_content
