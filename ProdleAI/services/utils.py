import io
import os
from docx import Document as docxDocument
from PyPDF2 import PdfReader
from spire.doc import *
from spire.doc.common import *
from spire.pdf import *
from spire.pdf.common import *
from pymongo import MongoClient
import pytesseract
from PIL import Image
from dotenv import load_dotenv
from openai import OpenAI as org_opeani
from openai import OpenAI
from anthropic import Anthropic
from ai_models import OpenaiModel, ClaudeModel
import requests
from prompt_templates.templates import entity_extraction
import tempfile, oci, traceback

import base64
from datetime import datetime, timedelta
    
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
tesseract_path = os.getenv("TESSERACT_PATH")
folder_path = os.getenv("FOLDER_PATH")
server_url = os.getenv("SERVER_URL")
mongo_client = MongoClient(os.getenv("MONGO_DB_URL"))

# Access your database
mongo_db = mongo_client["prodle_db"]
config = oci.config.from_file()  # or use from_environment()
object_storage_client = oci.object_storage.ObjectStorageClient(config)
namespace = object_storage_client.get_namespace().data
bucket_name = os.getenv("OCI_BUCKET_NAME")
region = config["region"]

def download_file_from_oci(object_name, org_id=None):
    """
    Download a file from OCI object storage.
    
    Args:
        object_name: The name of the object in OCI
        org_id: Optional organization ID to get specific OCI config
        
    Returns:
        Path to the downloaded temporary file
    """
    # Extract org_id from object_name if not provided
    if not org_id and '/' in object_name:
        parts = object_name.split('/')
        if len(parts) > 1:
            potential_org_id = parts[0]
            # Check if this is a valid org_id in the objectstores collection
            objectstore_collection = mongo_db["objectstores"]
            if objectstore_collection.find_one({"organizationId": potential_org_id}):
                org_id = potential_org_id
    
    # Get OCI configuration
    oci_config_raw = get_oci_config_from_objectstore(org_id)
    
    # Use the utility function to decode OCI config
    oci_config, namespace, bucket_name, key_file_path = decode_oci_config(oci_config_raw)
    
    # Create the OCI client
    object_storage_client = oci.object_storage.ObjectStorageClient(oci_config)
    
    try:
        # Get the object
        response = object_storage_client.get_object(
            namespace_name=namespace,
            bucket_name=bucket_name,
            object_name=object_name
        )
        # print("response in download_file_from_oci", response)
        
        # Save to a temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(object_name)[1])
        temp_file.write(response.data.content)
        temp_file.close()
        return temp_file.name  # path to use with docxDocument() or others
    finally:
        # Clean up the temporary key file
        if key_file_path and os.path.exists(key_file_path):
            try:
                os.unlink(key_file_path)
            except Exception as e:
                print(f"Error removing temporary key file: {str(e)}")


def generate_openai_embeddings(text_chunks):
    """
    Generate embeddings for a list of text chunks using OpenAI's Ada model.
    """
    embeddings = []
    openai_client = OpenAI()
    #print("text_chunks in generate openai embeddings", text_chunks)
    #print("length of text_chunks", len(text_chunks))
    try:
        for chunk in text_chunks:
            response = openai_client.embeddings.create(
                input=chunk,
                model="text-embedding-ada-002"
            )
            # Correctly access embedding data using response's structure
            embedding = response.data[0].embedding
            embeddings.append(embedding)
    except Exception as e:
        print(f"Error generating embeddings: {e}")
    return embeddings




pytesseract.pytesseract.tesseract_cmd = tesseract_path
# Function to extract text from paragraphs
def extract_text_from_paragraph(paragraph):
    return paragraph.text

def convert_table_data_to_string(table_data):
    #print(table_data)
    import json
    table_string = []
    for row in table_data:
        row_dict = {str(i): val for i, val in enumerate(row)}
        table_string.append(json.dumps(row_dict))
    return ' --END-- \n '.join(table_string)

def convert_df_to_string_json(table_data):
    print(table_data)
    table_string = ""
    for row in table_data:
        row_string = ", ".join([f'"{col}": "{val}"' for col, val in enumerate(row)])
        table_string += f"{{{row_string}}} --END-- \n"
    return table_string

def convert_df_to_string_kv(table_data):
    table_string = ""
    for row in table_data:
        row_string = " | ".join([f"{col}: {val}" for col, val in enumerate(row)])
        table_string += f"{row_string} --END-- \n"
    return table_string

def convert_df_to_string_xml(table_data):
    table_string = ""
    for row in table_data:
        row_string = " ".join([f"<{col}>{val}</{col}>" for col, val in enumerate(row)])
        table_string += f"<row> {row_string} </row> --END-- \n"
    return table_string

# Function to extract text from tables
def extract_text_from_table(table):
    table_data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        table_data.append(row_data)
    return table_data

# Function to extract images and apply OCR
def extract_images_and_ocr(doc):
    header = doc.sections[0].header
    footer = doc.sections[0].footer
    header_text = []
    footer_text = []
    for rel in header.part.rels:
        if "image" in header.part.rels[rel].target_ref:
            img = header.part.rels[rel].target_part.blob
            image = Image.open(io.BytesIO(img))
            text = pytesseract.image_to_string(image)
            header_text.append(text)
    for rel in footer.part.rels:
        if "image" in footer.part.rels[rel].target_ref:
            img = footer.part.rels[rel].target_part.blob
            image = Image.open(io.BytesIO(img))
            text = pytesseract.image_to_string(image)
            footer_text.append(text)
    images_text = {
        "header": header_text,
        "footer": footer_text
    }
    return images_text

# Main function to extract content from DOCX
def preprocess_docx(file_path):
    document = docxDocument(file_path)
    # images_text = extract_images_and_ocr(document)
    # header_image = images_text["header"]
    # footer_image = images_text["footer"]
    hearder_table = ''
    footer_table = ''
    
    for table in document.sections[0].header.tables:
        hearder_table = extract_text_from_table(table)

    for table in document.sections[0].footer.tables:
        footer_table = extract_text_from_table(table)
    
    
    headers = repr(hearder_table)
    # repr(header_image) 
    # + 
    
    footers = repr(footer_table)
    
    header_footer_content = headers + footers
    
    return header_footer_content

def get_entities(doc_text):
    openai_client = org_opeani(api_key=os.getenv("OPENAI_API_KEY"))
    text =  doc_text
    response = openai_client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": entity_extraction},
            {"role": "user", "content": text}

        ],
        temperature=0,
        max_tokens=4091,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response.choices[0].message.content

def get_entities_for_bulk_upload(doc_text, form_data):
    openai_client = org_opeani(api_key=os.getenv("OPENAI_API_KEY"))
    docTypeOptions = [value for key, value in form_data.items() if key.startswith('docTypeOptions')]
    prompt = entity_extraction.format(
        context_str = doc_text,
        docTypeOptions = docTypeOptions
    )
    response = openai_client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}

        ],
        temperature=0,
        max_tokens=4091,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response.choices[0].message.content

def remove_headers_footers(doc):
    sections = doc.sections
    for section in sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True


def get_embedding_for_text(text):
    """
    Generate embedding for input text using OpenAI's embedding model.
    """
    openai_client = OpenAI()
    response = openai_client.embeddings.create(
        input=text,
        model="text-embedding-ada-002"
    )
    return response.data[0].embedding

def calculate_gpt4_turbo_cost(usage):
    # Pricing for GPT-4o
    cost_per_million_input_tokens = 2.50  # $2.50 per 1M input tokens
    cost_per_million_output_tokens = 10.00  # $10.00 per 1M output tokens

    # Extracting token usage details
    prompt_tokens = usage.prompt_tokens
    completion_tokens = usage.completion_tokens

    # Calculating costs
    input_cost = (prompt_tokens / 1_000_000) * cost_per_million_input_tokens
    output_cost = (completion_tokens / 1_000_000) * cost_per_million_output_tokens

    # Total cost
    total_cost = input_cost + output_cost

    # Printing the usage details and cost
    print(f"Prompt Tokens (Input): {prompt_tokens}")
    print(f"Completion Tokens (Output): {completion_tokens}")
    print(f"Total Tokens: {usage.total_tokens}")
    print(f"Input Cost: ${input_cost:.4f}")
    print(f"Output Cost: ${output_cost:.4f}")
    print(f"Total Cost: ${total_cost:.4f}")

    return total_cost


def get_feature_config(org_record, featureName):
    """Fetches the summary configuration from the organization record."""
    return org_record["aiConfig"]['featureConfigs'][featureName]

def get_ai_llm_instance(org_record, client):
    """Fetches and initializes the LLM instance based on the client type."""
    api_keys = org_record["aiConfig"]['apiKeys']
    model = org_record["aiConfig"]['featureConfigs']['Chat']['model']

    if client == "OpenAI":
        api_key = api_keys["OpenAI"]
        # Instantiate the OpenAI LLM with the specified parameters
        llm = OpenAI(temperature=0, model=model, api_key=api_key)
        return llm

    elif client == "Anthropic":
        api_key = api_keys["Anthropic"]
        # Instantiate the Anthropic client
        client = Anthropic(api_key=api_key)
        return client

    else:
        raise ValueError(f"Unsupported client type: {client}")

def get_ai_model_config(org_record, client):
    """Fetches the API key and model class based on the client type."""
    # print("client in get ai model config", client)
    api_keys = org_record["aiConfig"]['apiKeys']
    if client == "OpenAI":
        return api_keys["OpenAI"], OpenaiModel
    elif client == "Anthropic":
        # print("api keys", api_keys["Anthropic"])
        return api_keys["Anthropic"], ClaudeModel
    else:
        raise ValueError(f"Unsupported client type: {client}")

def setup_model(api_key, model_class, model_name, max_tokens=4091, temperature=0):
    """Sets up the AI model with the given configuration."""
    return model_class(config={
        "api_key": api_key,
        "model": model_name,
        "max_tokens": max_tokens,
        "temperature": temperature
    })

def docx_to_text_using_spire(file_path):
    # Create a Document object
    document = Document()
    # Load a Word document
    document.LoadFromFile(file_path)

    # Extract the text of the document
    document_text = document.GetText()

    return document_text


def pdf_to_text(file_path):
    reader = PdfReader(file_path)
    output = []
    for page in reader.pages:
        text = page.extract_text()
        output.append(text)

    return "\n".join(output)


def file_to_text(file_path):
    try:
        if file_path.endswith(".docx"):
            text = docx_to_text_using_spire(file_path)
            # print("text using spire----------", text)
            # text_docx = docx_to_text(file_path)
            # print("text using docx----------", text_docx)
            return text
        elif file_path.endswith(".pdf"):
            doc = PdfDocument()
            doc.LoadFromFile(file_path)

            extracted_text = []

            # Iterate over the pages of the document
            for i in range(doc.Pages.Count):
                    page = doc.Pages.get_Item(i)
                    # Extract the text from the page
                    textExtractor = PdfTextExtractor(page)
                    option = PdfTextExtractOptions()        
                    text = textExtractor.ExtractText(option)
                    extracted_text.append(text)

            # print("extracted_text", extracted_text)

            doc.Close()
            # Join the extracted text from all pages into one complete string
            return " ".join(extracted_text)
        else:
            raise ValueError("Unsupported file format")
    except Exception as e:
        print("Error:", e)
        return None


def get_org_details_by_id(org_id):
    org_collection = mongo_db["Organization"]
    org_record = org_collection.find_one({"_id" : org_id})
    return org_record;

def insert_into_ai_meta_data_collection(record):
    ai_meta_data_collection = mongo_db["aimetadatas"]
    print("record in insert_into_ai_meta_data_collection", ai_meta_data_collection)
    ai_meta_data_collection.insert_one(record)

def insert_into_audit_checklist_collection(record):
    audit_dept_checklist_collection = mongo_db["auditdeptchecklists"]
    print("record in audit_dept_checklist_collection", audit_dept_checklist_collection)
    audit_dept_checklist_collection.insert_one(record)

def call_create_schedule_api(schedule_data):
    nestjs_url = f"{server_url}/api/auditSchedule/createAuditScheduleForPython"
    headers = {
        'Content-Type': 'application/json'
    }
    response = requests.post(nestjs_url, headers=headers, json=schedule_data)

def insert_into_question_collection(question_data):
    collection = mongo_db["questions"]
    print("question_data in insert_into_question_collection", question_data)
    result = collection.insert_one(question_data)
    print("result plain", result)
    print("result json", result.inserted_id)
    print("Inserted question with id:", result.inserted_id)
    return result.inserted_id

def insert_into_audit_template_collection(template_data):
    collection = mongo_db["audittemplates"]
    print("template_data in insert_into_audit_template_collection", template_data)
    result = collection.insert_one(template_data)
    print("result of insert into audit template collection", result)
    template_id = str(result.inserted_id)  # Convert ObjectId to string
    print("Inserted template with id:", template_id)
    return template_id


def get_hira_data(hira_id):
    nestjs_url = f"/api/riskregister/hira/getHiraWithStepsWithoutToken/{hira_id}?&page=1&pageSize=200"
    headers = {
        'Content-Type': 'application/json'
    }
    response = requests.get(f"{server_url}{nestjs_url}", headers=headers)
    #print("response from hira get data nestjs", response.json())
    return response.json()

def get_capa_data(capa_id):
    nestjs_url = f"/api/cara/getCapaById/{capa_id}"
    headers = {
        'Content-Type': 'application/json'
    }
    try:
        response = requests.get(f"{server_url}{nestjs_url}", headers=headers, timeout=5)
        response.raise_for_status()  # Raises an HTTPError for bad responses (4xx, 5xx)

        if not response.text.strip():  # Handle empty responses
            return None

        return response.json()

    except requests.exceptions.JSONDecodeError:
        print(f"Error: Invalid JSON response for ID {capa_id}")
    except requests.exceptions.RequestException as e:
        print(f"Request failed: {e}")

    return None
    #print("response from capa get data nestjs", response.json())
    return response.json()



def call_create_schedule_api(schedule_data):
    """Calls the NestJS API to create an audit schedule and returns the schedule ID."""
    nestjs_url = f"{server_url}/api/auditSchedule/createAuditScheduleForPython"
    headers = {
        'Content-Type': 'application/json'
    }

    response = requests.post(nestjs_url, headers=headers, json=schedule_data)
    
    
    if response.status_code in [200, 201]:  
        response_data = response.json()
        print("Schedule API Response:", response_data)
        
        schedule_id = response_data.get("_id") 
        if schedule_id:
            print(f" Schedule successfully created with ID: {schedule_id}")
            return schedule_id
        else:
            print("Failed to extract schedule ID from response.")
            return None
    else:
        print(f" Failed to create schedule. Status Code: {response.status_code}, Response: {response.text}")
        return None


def get_all_risk_categories(org_id):
    config_collection = mongo_db["hiraconfigs"]
    org_records = config_collection.find({"organizationId": org_id})
    
    # Extract riskCategory from each document
    risk_categories = [record["riskCategory"] for record in org_records if "riskCategory" in record]
    
    return risk_categories

def get_oci_config_from_objectstore(org_id=None):
    """
    Get OCI configuration from the objectstores collection.
    If org_id is provided, it will try to find the configuration for that organization.
    If not found or org_id is None, it will fall back to the 'master' configuration.
    
    Returns a dictionary with OCI configuration parameters.
    """
    objectstore_collection = mongo_db["objectstores"]
    
    # Try to find configuration for the specific organization
    if org_id:
        config = objectstore_collection.find_one({"organizationId": org_id})
        if config:
            return config
    
    # Fall back to master configuration
    config = objectstore_collection.find_one({"organizationId": "master"})
    if not config:
        raise Exception("No OCI configuration found in objectstores collection")
    
    return config


def decode_oci_config(oci_config_raw):
    """
    Decode base64 encoded OCI configuration values and return a properly formatted config object.
    
    Args:
        oci_config_raw (dict): Raw OCI configuration with base64 encoded values
        
    Returns:
        tuple: (oci_config, namespace, bucket_name, private_key)
            - oci_config (dict): Formatted OCI configuration
            - namespace (str): Decoded namespace
            - bucket_name (str): Decoded bucket name
            - private_key (str): Private key
    """
    try:
        # Decode base64 encoded values
        user_id = base64.b64decode(oci_config_raw.get("userId")).decode('utf-8')
        tenancy_id = base64.b64decode(oci_config_raw.get("tenancyId")).decode('utf-8')
        fingerprint = base64.b64decode(oci_config_raw.get("fingerprint")).decode('utf-8')
        region = base64.b64decode(oci_config_raw.get("region")).decode('utf-8')
        namespace = base64.b64decode(oci_config_raw.get("namespace")).decode('utf-8')
        bucket_name = base64.b64decode(oci_config_raw.get("bucketName")).decode('utf-8')
        private_key = oci_config_raw.get("privateKey")
        
        # Create a temporary file for the private key
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pem') as temp_key_file:
            temp_key_file.write(private_key)
            key_file_path = temp_key_file.name
        
        # Create a properly formatted OCI config
        oci_config = {
            "user": user_id,
            "tenancy": tenancy_id,
            "fingerprint": fingerprint,
            "key_file": key_file_path,
            "region": region
        }
        
        return oci_config, namespace, bucket_name, key_file_path
    except Exception as e:
        print(f"Error decoding OCI config: {str(e)}")
        raise

def generate_preauthenticated_url(object_name, org_id=None, expiry_seconds=300):
    """
    Generate a preauthenticated URL for an OCI object.
    
    Args:
        object_name: The name of the object in OCI
        org_id: Optional organization ID to get specific OCI config
        expiry_seconds: Number of seconds until the URL expires (default: 300)
        
    Returns:
        A preauthenticated URL for the object
    """
    
    # Get OCI configuration from objectstores collection
    config = get_oci_config_from_objectstore(org_id)
    
    # Use the utility function to decode OCI config
    oci_config, namespace, bucket_name, key_file_path = decode_oci_config(config)
    
    # Create OCI client
    object_storage_client = oci.object_storage.ObjectStorageClient(oci_config)
    
    # Create preauthenticated request details
    expiry_time = datetime.now() + timedelta(seconds=expiry_seconds)
    create_preauthenticated_request_details = oci.object_storage.models.CreatePreauthenticatedRequestDetails(
        name=f"par-object-{datetime.now().strftime('%Y%m%d%H%M%S')}",
        object_name=object_name,
        access_type="AnyObjectRead",
        time_expires=expiry_time
    )
    
    # Create preauthenticated request
    create_preauthenticated_request_request = oci.object_storage.requests.CreatePreauthenticatedRequestRequest(
        namespace_name=namespace,
        bucket_name=bucket_name,
        create_preauthenticated_request_details=create_preauthenticated_request_details
    )
    
    # Get preauthenticated request response
    response = object_storage_client.create_preauthenticated_request(create_preauthenticated_request_request)
    
    # Clean up the temporary key file
    if key_file_path and os.path.exists(key_file_path):
        try:
            os.unlink(key_file_path)
        except Exception as e:
            print(f"Error removing temporary key file: {str(e)}")
    
    # Return the full preauthenticated URL
    return response.data.full_path + object_name
