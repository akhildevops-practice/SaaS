import docx
import pytesseract
from PIL import Image
import io
import os
from services.utils import get_capa_data,generate_openai_embeddings
from prompt_templates.templates import entity_extraction,capa_ai_suggestions_prompt
from openai import OpenAI as org_opeani
from openai import OpenAI
from dotenv import load_dotenv
from ai_models import OpenaiModel, ClaudeModel
from flask import Flask, jsonify
from flask_cors import CORS
from pymilvus import Collection
from database.connection import connect_to_milvus, use_database
from database.settings import DEFAULT_DB_NAME
import json
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
tesseract_path = os.getenv("TESSERACT_PATH")
folder_path = os.getenv("FOLDER_PATH")
server_url = os.getenv("SERVER_URL")
app = Flask(__name__)
api_key = os.getenv("OPENAI_API_KEY")
CORS(app, origins="*")
pytesseract.pytesseract.tesseract_cmd = tesseract_path
from sentence_transformers import CrossEncoder

# Load the Cross-Encoder model (pre-trained for MS MARCO)
cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2')
# Function to extract text from paragraphs
def extract_text_from_paragraph(paragraph):
    return paragraph.text

def convert_table_data_to_string(table_data):
    #print(table_data)
    import json
    table_string = []
    for row in table_data:
        row_dict = {str(i): val for i, val in enumerate(row)}
        table_string.append(json.dumps(row_dict))
    return ' --END-- \n '.join(table_string)

def convert_df_to_string_json(table_data):
    table_string = ""
    for row in table_data:
        row_string = ", ".join([f'"{col}": "{val}"' for col, val in enumerate(row)])
        table_string += f"{{{row_string}}} --END-- \n"
    return table_string

def convert_df_to_string_kv(table_data):
    table_string = ""
    for row in table_data:
        row_string = " | ".join([f"{col}: {val}" for col, val in enumerate(row)])
        table_string += f"{row_string} --END-- \n"
    return table_string

def convert_df_to_string_xml(table_data):
    table_string = ""
    for row in table_data:
        row_string = " ".join([f"<{col}>{val}</{col}>" for col, val in enumerate(row)])
        table_string += f"<row> {row_string} </row> --END-- \n"
    return table_string

# Function to extract text from tables
def extract_text_from_table(table):
    table_data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        table_data.append(row_data)
    return table_data

# Function to extract images and apply OCR
def extract_images_and_ocr(doc):
    header = doc.sections[0].header
    footer = doc.sections[0].footer
    header_text = []
    footer_text = []
    for rel in header.part.rels:
        if "image" in header.part.rels[rel].target_ref:
            img = header.part.rels[rel].target_part.blob
            image = Image.open(io.BytesIO(img))
            text = pytesseract.image_to_string(image)
            header_text.append(text)
    for rel in footer.part.rels:
        if "image" in footer.part.rels[rel].target_ref:
            img = footer.part.rels[rel].target_part.blob
            image = Image.open(io.BytesIO(img))
            text = pytesseract.image_to_string(image)
            footer_text.append(text)
    images_text = {
        "header": header_text,
        "footer": footer_text
    }
    return images_text

# Main function to extract content from DOCX
def preprocess_docx(file_path):
    document = docx.Document(file_path)
    # images_text = extract_images_and_ocr(document)
    # header_image = images_text["header"]
    # footer_image = images_text["footer"]
    hearder_table = ''
    footer_table = ''
    
    text_data = []
    bullet_points = []
    tables = []
    
    # for paragraph in document.paragraphs:
    #     if paragraph.style.name.startswith('Heading'):
    #         text_data.append(paragraph.text)
    #     elif paragraph.style.name in ['ListBullet', 'ListNumber']:
    #         bullet_points.append(paragraph.text)
    #     else:
    #         text_data.append(paragraph.text)
    
    for table in document.sections[0].header.tables:
        hearder_table = extract_text_from_table(table)

    for table in document.sections[0].footer.tables:
        footer_table = extract_text_from_table(table)
    
    # for table in document.tables:
    #     tables.append(extract_text_from_table(table))

    headers = repr(hearder_table)
    # repr(header_image) 
    # + 
    
    footers = repr(footer_table)
    
    header_footer_content = headers + footers
    
    return header_footer_content

def get_entities(doc_text):
    openai_client = org_opeani(api_key=os.getenv("OPENAI_API_KEY"))
    text =  doc_text
    response = openai_client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": entity_extraction},
            {"role": "user", "content": text}

        ],
        temperature=0,
        max_tokens=4091,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response.choices[0].message.content

def get_entities_for_bulk_upload(doc_text, form_data):
    openai_client = org_opeani(api_key=os.getenv("OPENAI_API_KEY"))
    docTypeOptions = [value for key, value in form_data.items() if key.startswith('docTypeOptions')]
    prompt = entity_extraction.format(
        context_str = doc_text,
        docTypeOptions = docTypeOptions
    )
    response = openai_client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}

        ],
        temperature=0,
        max_tokens=4091,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response.choices[0].message.content

def remove_headers_footers(doc):
    sections = doc.sections
    for section in sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True


def get_embedding_for_text(text):
    """
    Generate embedding for input text using OpenAI's embedding model.
    """
    openai_client = OpenAI()
    response = openai_client.embeddings.create(
        input=text,
        model="text-embedding-ada-002"
    )
    return response.data[0].embedding


def get_capa_suggestions(body):
    """
    Fetch CAPA suggestions based on query using Milvus.
    """
    connect_to_milvus()
    use_database(DEFAULT_DB_NAME)
    collection_name = "capa_collection"
    collection = Collection(name=collection_name)
    # Ensure the collection has an index; if not, create one
    if not collection.has_index():
        index_params = {
            "index_type": "IVF_FLAT",  # Choose index type suitable for your use case
            "metric_type": "IP",       # Inner Product for cosine similarity
            "params": {"nlist": 128}
        }
        collection.create_index(field_name="embedding", index_params=index_params)

    # Load the collection into memory
    collection.load()
    
    # Construct CAPA text from various fields
    CAPA_text = f"Problem Statement: {body.get('title', 'N/A')}\n"
    CAPA_text += f"Description: {body.get('description', 'N/A')}\n"
    CAPA_text += f"Organization: {body.get('organizationName', 'N/A')}\n"
    CAPA_text += f"Origin: {body.get('origin', 'N/A')}\n"
    CAPA_text += f"registeredBy: {body.get('origin', 'N/A')}\n"
    CAPA_text += f"Location: {body.get('locationId', 'N/A')}\n"
    CAPA_text += f"Entity: {body.get('entityId', 'N/A')}\n"
    CAPA_text += f"Containment Action: {body.get('containmentAction', 'N/A')}\n"
    CAPA_text += f"Why 5: {body.get('why5', 'N/A')}\n"  # Using .get() to avoid KeyError
    CAPA_text += f"Why 4: {body.get('why4', 'N/A')}\n"
    CAPA_text += f"Why 3: {body.get('why3', 'N/A')}\n"
    CAPA_text += f"Why 2: {body.get('why2', 'N/A')}\n"
    CAPA_text += f"Why 1: {body.get('why1', 'N/A')}\n"
    CAPA_text += f"Method: {body.get('method', 'N/A')}\n"
    CAPA_text += f"Measurement: {body.get('measurement', 'N/A')}\n"
    CAPA_text += f"Material: {body.get('material', 'N/A')}\n"
    CAPA_text += f"Man: {body.get('man', 'N/A')}\n"
    CAPA_text += f"Machine: {body.get('machine', 'N/A')}\n"
    CAPA_text += f"Environment: {body.get('environment', 'N/A')}\n"
    CAPA_text += f"Corrective Action: {body.get('correctiveAction', 'N/A')}\n"
    CAPA_text += f"Root Cause Analysis: {body.get('rootCauseAnalysis', 'N/A')}\n"
    CAPA_text += f"Defect Type: {body.get('defectType', 'N/A')}\n"
    CAPA_text += f"Advanced Root Cause: {body.get('advancedRootCause', 'N/A')}\n"
    CAPA_text += f"Fishbone: {body.get('fishBone', 'N/A')}\n"
    CAPA_text += f"Is-Is Not: {body.get('isIsNot', 'N/A')}\n"
    CAPA_text += f"Outcome: {body.get('outcome', 'N/A')}\n"
    embedding_for_query = get_embedding_for_text(CAPA_text)

    search_params = {
        "metric_type": "IP",
        "params": {"nprobe": 10}
    }
    results = collection.search(
        data=[embedding_for_query],
        anns_field="embedding",
        param=search_params,
        limit=50,
        expr=None,
    )

    suggestions = []
    query_doc_pairs = []
    seen_capa_ids = set()
    for result in results[0]:  # Loop through the first query's results
        
        doc_metadata = collection.query(
            expr=f"id == {result.id}",
            output_fields=["capa_id", "capa_text"]
        )[0]
        
        capa_id = doc_metadata["capa_id"]
        if capa_id in seen_capa_ids:
            continue  # Skip if already seen
        
        seen_capa_ids.add(capa_id)
        capa_data = get_capa_data(doc_metadata["capa_id"])
        if not capa_data or "content" not in capa_data:
            print(f"Warning: No valid CAPA content found for ID {capa_id}")
            continue  # Skip if the response is empty or missing 'content'
        CAPA = capa_data["content"]
        # Construct CAPA text from various fields
        CAPA_text_nestjs = f"Problem Statement: {CAPA.get('title', 'N/A')}\n"
        CAPA_text_nestjs = f"Description: {CAPA.get('description', 'N/A')}\n"
        CAPA_text_nestjs += f"Location: {CAPA.get('locationId', 'N/A')}\n"
        CAPA_text_nestjs += f"Entity: {CAPA.get('entityId', 'N/A')}\n"
        CAPA_text_nestjs += f"Containment Action: {CAPA.get('containmentAction', 'N/A')}\n"
        CAPA_text_nestjs += f"Why 5: {CAPA.get('why5', 'N/A')}\n"  # Using .get() to avoid KeyError
        CAPA_text_nestjs += f"Why 4: {CAPA.get('why4', 'N/A')}\n"
        CAPA_text_nestjs += f"Why 3: {CAPA.get('why3', 'N/A')}\n"
        CAPA_text_nestjs += f"Why 2: {CAPA.get('why2', 'N/A')}\n"
        CAPA_text_nestjs += f"Why 1: {CAPA.get('why1', 'N/A')}\n"
        CAPA_text_nestjs += f"Method: {CAPA.get('method', 'N/A')}\n"
        CAPA_text_nestjs += f"Measurement: {CAPA.get('measurement', 'N/A')}\n"
        CAPA_text_nestjs += f"Material: {CAPA.get('material', 'N/A')}\n"
        CAPA_text_nestjs += f"Man: {CAPA.get('man', 'N/A')}\n"
        CAPA_text_nestjs += f"Machine: {CAPA.get('machine', 'N/A')}\n"
        CAPA_text_nestjs += f"Environment: {CAPA.get('environment', 'N/A')}\n"
        CAPA_text_nestjs += f"Corrective Action: {CAPA.get('correctiveAction', 'N/A')}\n"
        CAPA_text_nestjs += f"Root Cause Analysis: {CAPA.get('rootCauseAnalysis', 'N/A')}\n"
        CAPA_text_nestjs += f"Advanced Root Cause Analysis: {CAPA.get('advancedRootCause', 'N/A')}\n"
        CAPA_text_nestjs += f"Fishbone: {CAPA.get('fishbone', 'N/A')}\n"
        CAPA_text_nestjs += f"Outcome: {CAPA.get('outcome', 'N/A')}\n"
        CAPA_text_nestjs += f"Is-Is Not: {CAPA.get('isIsNot', 'N/A')}\n"
        
        similarity_score = result.distance
        metadata_score = compute_metadata_score(CAPA)

        combined_score = (
            0.5 * similarity_score + 
            0.3 * metadata_score 
        )

        k = 60
        rrf_score = 1 / (k + (len(suggestions) + 1))
        
        query_doc_pairs.append((CAPA_text, CAPA_text_nestjs))
        
        # Append the suggestion
        suggestions.append({
            "text": doc_metadata["capa_text"],
            "capa_text_full": CAPA_text_nestjs,
            "capa_data": CAPA,
            "combined_score": combined_score,
            "metadata_score": metadata_score,
            "similarity_score": similarity_score,
            "rrf_score": rrf_score
        })
        
    reRankingType = body.get('reRankingType', 1)

    if reRankingType == 1:
        # Weighted Ranker
        for suggestion in suggestions:
            suggestion["final_score"] = suggestion["combined_score"]
    elif reRankingType == 2:
        # RRF Ranker
        for suggestion in suggestions:
            suggestion["final_score"] = suggestion["rrf_score"]
    elif reRankingType == 3:
        #Cross Encoding
        if query_doc_pairs:
            scores = cross_encoder.predict(query_doc_pairs)
            for i, suggestion in enumerate(suggestions):
                suggestion["final_score"] = scores[i]
            

    # **Manual reranking based on final_score**
    suggestions = sorted(suggestions, key=lambda x: x["final_score"], reverse=True)

    return suggestions[:5] if suggestions else None

def compute_metadata_score(capa_data):
    score = 0
    if capa_data.get("containmentAction"): score += 0.2
    if capa_data.get("correctiveAction"): score += 0.2
    if capa_data.get("rootCauseAnalysis"): score += 0.3
    if capa_data.get("method"): score += 0.1
    if capa_data.get("measurement"): score += 0.1
    if capa_data.get("outcome"): score += 0.2
    if capa_data.get("advancedRootCause"): score += 0.1
    if capa_data.get("fishbone"): score += 0.1
    return score


def search_capa_in_doc(body):
    """
    Search documents in Milvus using a query embedding.
    """
    connect_to_milvus()
    use_database(DEFAULT_DB_NAME)
    collection_name = "capa_collection"
    collection = Collection(name=collection_name)
    # Ensure the collection has an index; if not, create one
    if not collection.has_index():
        index_params = {
            "index_type": "IVF_FLAT",  # Choose index type suitable for your use case
            "metric_type": "IP",       # Inner Product for cosine similarity
            "params": {"nlist": 128}
        }
        collection.create_index(field_name="embedding", index_params=index_params)

    # Load the collection into memory
    collection.load()
    

    query_text = body["query"]
    embedding_for_query = get_embedding_for_text(query_text)

    search_params = {
        "metric_type": "IP",
        "params": {"nprobe": 10}
    }
    results = collection.search(
        data=[embedding_for_query],
        anns_field="embedding",
        param=search_params,
        limit=body.get("top_k", 4),
        expr=None,
    )

    search_results = []
    for result in results[0]:
        doc_metadata = collection.query(
            expr=f"id == {result.id}",
            output_fields=["capa_id", "capa_text"]
        )[0]
        search_results.append({
            "text": doc_metadata["capa_text"],
            "similarity": result.distance,
            "capaId": doc_metadata["capa_id"],
        })

    return search_results
def chunk_text(text, max_length=3000):
    chunks = []
    while len(text) > max_length:
        split_point = text.rfind("\n", 0, max_length)
        if split_point == -1:
            split_point = max_length
        chunks.append(text[:split_point])
        text = text[split_point:].strip()
    chunks.append(text)
    return chunks
def generate_embeddings_and_prepare_data(CAPA_text, metadata):
    text_chunks = chunk_text(CAPA_text)
    #print(f"Number of text chunks: {text_chunks}")
    embeddings = generate_openai_embeddings(text_chunks)
    #print(f"Number of text chunks: {len(embeddings)}")

    # if len(embeddings) != len(text_chunks):
    #     raise ValueError("Number of embeddings does not match number of text chunks.")

    # Prepare data for insertion into Milvus
    chunk_ids = list(range(len(text_chunks)))
    metadata_list = [{**metadata, "chunkIndex": i} for i in range(len(text_chunks))]

    return text_chunks, embeddings, chunk_ids, metadata_list


def insert_capa_in_milvus_capa_collection(capa_id):
    capa_data = get_capa_data(capa_id)
    # Use the content as a dictionary (no need for json.dumps)
    content = capa_data["content"]
    CAPA = capa_data["content"]
    # Extract meta_data correctly (content is already a dictionary)
    meta_data = {
        "organizationId": content.get("organizationId", "N/A"),
        "capaId": content.get("_id", "N/A"),
        "title": content.get("title", "N/A"),
        "kpiId": content.get("kpiId", "N/A"),
        "registeredBy": content.get("registeredBy", {}).get("id", "N/A"),
        "caraCoordinator": content.get("caraCoordinator", {}).get("id", "N/A"),
        "caraOwner": content.get("caraOwner", {}).get("id", "N/A"),
        "location": content.get("locationId", "N/A"),
        "entity": content.get("entityId", {}).get("id", "N/A"),
        "systems": list(map(lambda x: x.get("_id", "N/A"), content.get("systemId", []))),
        "description": content.get("description", "N/A"),
        "origin": content.get("origin", {}).get("_id", "N/A"),
        "correctedDate": content.get("correctedDate", "N/A"),
        "type": content.get("type", "N/A"),
        "serialNumber": content.get("serialNumber", "N/A"),
        "highPriority": content.get("highPriority", "N/A"),
        "defectType": content.get("defectType", "N/A"),
            }
    capa_details = {
        "organizationName": content.get("organizationName", "N/A"),
        "location": content.get("locationDetails", {}).get("locationName", "N/A"),
        "entity": content.get("entityId", {}).get("entityName", "N/A"),
        "title": content.get("title", "N/A"),
        "description": content.get("description", "N/A"),
        "origin": content.get("origin", {}).get("deviationType", "N/A"),
        "registeredBy": content.get("registeredBy", {}).get("firstname", "N/A") + " " + content.get("registeredBy", {}).get("lastname", "N/A"),
        "caraCoordinator": content.get("caraCoordinator", {}).get("firstname", "N/A") + " " + content.get("caraCoordinator", {}).get("lastname", "N/A"),
        "caraOwner": content.get("caraOwner", {}).get("name", "N/A"),
        "systems": ", ".join(map(lambda x: x.get("name", "Unknown"), content.get("systemId", []))),
        "actualCorrectiveAction": content.get("actualCorrectiveAction", "N/A"),
        "containmentAction": content.get("containmentAction", "N/A"),
        "rootCauseAnalysis": content.get("rootCauseAnalysis", "N/A"),
        "correctiveAction": content.get("correctiveAction", "N/A"),
        "correctedDate": content.get("correctedDate", "N/A"),
        "type": content.get("type", "N/A"),
        "serialNumber": content.get("serialNumber", "N/A"),
        "comments": content.get("comments", "N/A"),
        "why1": content.get("why1", "N/A"),
        "why2": content.get("why2", "N/A"),
        "why3": content.get("why3", "N/A"),
        "why4": content.get("why4", "N/A"),
        "why5": content.get("why5", "N/A"),
        "man": content.get("man", "N/A"),
        "material": content.get("material", "N/A"),
        "measurement": content.get("measurement", "N/A"),
        "method": content.get("method", "N/A"),
        "machine": content.get("machine", "N/A"),
        "highPriority": content.get("highPriority", "N/A"),
        "defectType": content.get("defectType", "N/A"),
        "advancedRootCause": content.get("advancedRootCause", "N/A"),
        "fishBone": content.get("fishBone", "N/A"),
        "isIsNot": content.get("isIsNot", "N/A"),
        "outcome": content.get("outcome", "N/A"),
    }

    
    naturalText = generate_natural_language_summary(capa_details)
   # print("capa data", capa_data)
    
    #print("CONTENT in insert_Capa_in_milvus ", content)

    # Connect to Milvus
    connect_to_milvus()
    use_database(DEFAULT_DB_NAME)
    collection_name = "capa_collection"
    collection = Collection(name=collection_name)
    
    # Ensure the collection has an index; if not, create one
    if not collection.has_index():
        index_params = {
            "index_type": "IVF_FLAT",  # Choose index type suitable for your use case
            "metric_type": "IP",       # Inner Product for cosine similarity
            "params": {"nlist": 128}
        }
        collection.create_index(field_name="embedding", index_params=index_params)

    # Load the collection into memory
    collection.load()
    
    #print("Meta data:", meta_data)
    
    # Construct CAPA text from various fields
    CAPA_text = f"Problem Statement: {CAPA.get('title', 'N/A')}\n"
    CAPA_text = f"Description: {CAPA.get('description', 'N/A')}\n"
    CAPA_text += f"Location: {CAPA.get('locationDetails', 'N/A')}\n"
    CAPA_text += f"Entity: {CAPA.get('entityId', 'N/A')}\n"
    CAPA_text += f"Containment Action: {CAPA.get('containmentAction', 'N/A')}\n"
    CAPA_text += f"Why 5: {CAPA.get('why5', 'N/A')}\n"  # Using .get() to avoid KeyError
    CAPA_text += f"Why 4: {CAPA.get('why4', 'N/A')}\n"
    CAPA_text += f"Why 3: {CAPA.get('why3', 'N/A')}\n"
    CAPA_text += f"Why 2: {CAPA.get('why2', 'N/A')}\n"
    CAPA_text += f"Why 1: {CAPA.get('why1', 'N/A')}\n"
    CAPA_text += f"Method: {CAPA.get('method', 'N/A')}\n"
    CAPA_text += f"Measurement: {CAPA.get('measurement', 'N/A')}\n"
    CAPA_text += f"Material: {CAPA.get('material', 'N/A')}\n"
    CAPA_text += f"Man: {CAPA.get('man', 'N/A')}\n"
    CAPA_text += f"Machine: {CAPA.get('machine', 'N/A')}\n"
    CAPA_text += f"Environment: {CAPA.get('environment', 'N/A')}\n"
    CAPA_text += f"Corrective Action: {CAPA.get('correctiveAction', 'N/A')}\n"
    CAPA_text += f"Root Cause Analysis: {CAPA.get('rootCauseAnalysis', 'N/A')}\n"
    CAPA_text += f"Advanced Root Cause: {CAPA.get('advancedRootCause', 'N/A')}\n"
    CAPA_text += f"Is-Is Not: {CAPA.get('isIsNot', 'N/A')}\n"
    CAPA_text += f"Fishbone Analysis: {CAPA.get('fishBone', 'N/A')}\n"
    CAPA_text += f"Advanced Corrective Action: {CAPA.get('outcome', 'N/A')}\n"
    CAPA_text += f"Defect type: {CAPA.get('defectType', 'N/A')}\n"
    # Generate embeddings and prepare data (assuming generate_embeddings_and_prepare_data works)
    text_chunks, embeddings, chunk_ids, metadata_list = generate_embeddings_and_prepare_data(naturalText, meta_data)
    #print("text chunks and chunk ids",embeddings)
    # Prepare the data for insertion
    
    data_to_insert = [
        [meta_data["capaId"]] * len(chunk_ids),
        embeddings,                              # Embeddings
        text_chunks,
        metadata_list                            # Text chunks
    ]
    # data_to_insert = [
    # ["CAPA-12345"], 
    # [[0.12] * 1536],  
    # ["This is a sample CAPA document for testing."],  
    # [{"location": "Mumbai", "entity": "ABC Corp"}]  
    # ]
    
    # Insert data into Milvus collection
    try:
        result = collection.insert(data_to_insert)
        #print("Chunks added:", result)
        return "Success"

    except Exception as e:
        print(f"Error inserting data into Milvus: {e}")
        return jsonify({"message": "Failed to upload data to Milvus", "error": str(e)}), 500

def generate_capa_ai_suggestions(capa_dtls):
    openai_client = org_opeani(api_key=os.getenv("OPENAI_API_KEY"))
    prompt = capa_ai_suggestions_prompt.format(
        problemStatement=capa_dtls["problemStatement"],
        problemDescription=capa_dtls["problemDescription"],
        containmentAction= capa_dtls["containmentAction"],
        rootCauseAnalysis= capa_dtls["rootCauseAnalysis"],
        correctiveAction= capa_dtls["correctiveAction"],
        environment= capa_dtls["environment"],
        machine= capa_dtls["machine"],
        man= capa_dtls["man"],
        material= capa_dtls["material"],
        measurement= capa_dtls["measurement"],
        method= capa_dtls["method"],
        why1= capa_dtls["why1"],
        why2= capa_dtls["why2"],
        why3= capa_dtls["why3"],
        why4= capa_dtls["why4"],
        why5= capa_dtls["why5"],
        preferedTechnique= capa_dtls["preferedTechnique"]
    )
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=4091,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        suggestions = response.choices[0].message.content
        #print("suggestions", suggestions)
        return suggestions
    except Exception as e:
        print("Error generating suggestions:", e)
        return None
def search_capa(body):
    """
    Query CAPA data using Milvus vector search and return a single consolidated result.
    """
    try:
        # Step 1: Generate query embedding
        query_embedding = get_embedding_for_text(body["query"])

        # Step 2: Connect to Milvus collection
        collection_name = "capa_collection"
        collection = Collection(name=collection_name)

        # Ensure the collection has an index; if not, create one
        if not collection.has_index():
            index_params = {
                "index_type": "IVF_FLAT",
                "metric_type": "IP",
                "params": {"nlist": 128}
            }
            collection.create_index(field_name="embedding", index_params=index_params)

        # Load the collection into memory
        collection.load()

        # Step 3: Prepare search parameters and dynamic filters
        top_k = 20  # Retrieve up to 5 matching results
        search_params = {
            "metric_type": "IP",  # Inner Product for cosine similarity
            "params": {"nprobe": 10}
        }

        # Dynamic filters for metadata fields
        filters = []
        metadata_keys = ["location", "entity", "organizationId","origin"]  # Metadata keys for filtering
        for key in metadata_keys:
            if key in body:
                filters.append(f"metadata['{key}'] == '{body[key]}'")

        filter_expr = " and ".join(filters) if filters else ""
        #print("filter_expr",filter_expr)
        # Step 4: Perform search in Milvus
        results = collection.search(
            data=[query_embedding],
            anns_field="embedding",
            param=search_params,
            limit=top_k,
            output_fields=["capa_text", "metadata", "embedding"],  # Return text chunk and metadata
            expr=filter_expr
        )
        
        # Step 4: Prepare candidate pairs for re-ranking
        candidates = [
            (body["query"], top_result.entity.get("capa_text"))
            for top_result in results[0]
            if top_result.entity.get("capa_text")  # Ensure text is available
        ]
        #print("searched capa",candidates)

        if not candidates:
            return {"answer": "No relevant CAPA records found.", "sources": []}

        # Step 5: Rank candidates using Cross-Encoder
        scores = cross_encoder.predict(candidates)

        # Step 6: Sort results by Cross-Encoder score
        ranked_results = sorted(
            zip(results[0], scores),
            key=lambda x: x[1],  # Sort by score (higher is better)
            reverse=True
        )
        
        #print("search results",results)
        # Step 5: Filter out invalid results
        filtered_results = [
            {
                "chunkText": result[0].entity.get("capa_text"),
                "metadata": result[0].entity.get("metadata"),
                "cross_score": result[1]
            }
            for result in ranked_results[:5]  # Keep only the best 5 results
        ]

        if not filtered_results:
            return {
                "answer": "No results found.",
                "highlight_text": "",
                "sources": []
            }

        # Step 6: Consolidate results for OpenAI prompt
        prompt_contexts = "\n\n".join(
            [
                f"Context {idx}:\n"
                f"Text: {result['chunkText']}\n"
                f"Metadata: {result['metadata']}\n"
                for idx, result in enumerate(filtered_results)
            ]
        )

        #print("prompt_contexts", prompt_contexts)

        chat_prompt = (
            """ 
            {prompt_contexts}
            Query: {query} 

            Please provide a comprehensive and detailed response to the query based strictly on the provided contexts. Your response should:

            1. Be detailed and explanatory, written in simple, easy-to-understand language
            2. Break down complex information into clear points
            3. Explain cause-and-effect relationships when present
            4. Include relevant details from the context that help build a complete understanding
            5. Use natural, conversational language while maintaining professionalism
            6. Provide examples or clarifications from the context when available

            Important Guidelines:
            - Strictly use ONLY information present in the provided contexts
            - Do not make assumptions or include external knowledge
            - If information is not available in the contexts, respond with: "The information is not available in the provided contexts."
            - Ensure the answer is comprehensive yet focused on the query

            Format the response as a JSON object with the following keys:
            {{
                "answer": A detailed, well-structured explanation that thoroughly addresses the query. Break down complex information into digestible parts and use clear language. Include relevant context and explanations that help build a complete understanding.
                 "metadata": {{
            "location": "Location of the issue",
            "entity": "Entity involved",
                    }},
        "sources": An array of objects containing:
                    "ProblemStatement":"Problem Statement of the issue",
                   "location": "Location of the issue",
                   "entity": "Entity involved",
                
              
            }}

            Ensure the response is:
            1. A valid JSON object
            2. Based solely on provided contexts
            3. Comprehensive and detailed
            4. Clear and easy to understand
            5. Well-structured and logical
            make sure the json is valid,
            Don't wrap the JSON in ```json ``` tags. It should be a plain JSON object.
            """
        )

        chat_prompt = chat_prompt.format(prompt_contexts=prompt_contexts, query=body["query"])

        # Step 7: Query OpenAI
        model = OpenaiModel(
            config={
                "api_key": os.getenv("OPENAI_API_KEY"),
                "model": "gpt-4o-mini",
                "max_tokens": 4000,
                "temperature": 0,
            }
        )
        response = model.generate_response(system_message=chat_prompt, messages=[body["query"]])
        #print("response from openai in search hira", response)
        response_text = response["output"]
        # Step 8: Parse OpenAI response
        try:
        # Ensure response is a proper string
            cleaned_response = response_text.strip()
            
            # Fix improperly formatted JSON by replacing control characters
            cleaned_response = cleaned_response.replace("\n", "").replace("\t", "").replace("\r", "")
            #print("cleaned response",cleaned_response)

            response_json = json.loads(cleaned_response)
            
        except json.JSONDecodeError:
            raise Exception("Invalid JSON response from OpenAI.")

        # Validate response structure
        # if not isinstance(response_json, dict) or "answer" not in response_json or "sources" not in response_json:
        #     raise Exception("Unexpected response structure from OpenAI.")

        # Ensure sources are in expected format
        # sources = response_json.get("sources", [])
        # for source in sources:
        #     if "Problem Statements" not in source or "capaId" not in source or "text" not in source:
        #         raise Exception("Invalid source structure in OpenAI response.")

        return response_json

    except Exception as e:
        raise Exception(f"Error in CAPA query: {e}")


def search_similar_capas_milvus(body):
    query = body.get("query", {})
    title = query.get("title", "")
    description = query.get("description", "")
    combined_text = f"{title}. {description}"
    
    query_embedding = get_embedding_for_text(combined_text)
    if not query_embedding:
        return []

    collection = Collection(name="capa_collection")
    if not collection.has_index():
        index_params = {
            "index_type": "IVF_FLAT",
            "metric_type": "IP",
            "params": {"nlist": 128}
        }
        collection.create_index(field_name="embedding", index_params=index_params)

    collection.load()

    # Get metadata filters from top-level body
    metadata_keys = ["location", "entity", "organizationId", "origin"]
    filters = [f"metadata['{key}'] == '{body[key]}'" for key in metadata_keys if key in body]
    filter_expr = " and ".join(filters)

    search_params = {
        "metric_type": "IP",
        "params": {"nprobe": 10}
    }

    results = collection.search(
        data=[query_embedding],
        anns_field="embedding",
        param=search_params,
        limit=20,
        output_fields=["capa_text", "metadata", "embedding"],
        expr=filter_expr
    )

    candidates = [
        (combined_text, hit.entity.get("capa_text"))
        for hit in results[0]
        if hit.entity.get("capa_text")
    ]
    #print("candidates", candidates)
    
    if not candidates:
        return {"answer": "No relevant CAPA records found.", "sources": []}

    final_k = 5

    scores = cross_encoder.predict(candidates)
    #print("Cross-encoder scores:", scores)

    ranked_results = sorted(
        zip(results[0], scores),
        key=lambda x: x[1],
        reverse=True
    )

    print("Ranked results:", ranked_results)

    filtered_results = [
        {
            "chunkText": res.entity.get("capa_text"),
            "metadata": res.entity.get("metadata"),
            "cross_score": float(f"{score * 100:.2f}")
        }
        for res, score in ranked_results[:final_k]
    ]

    #print("Filtered result", filtered_results)

    return {
        "results": filtered_results,
        "answer": "Top similar CAPAs retrieved successfully."
    }

   
def generate_natural_language_summary(capa):
    """Use GPT to convert the CAPA object into a natural language summary."""

    # Check presence of advanced analysis data
    has_advanced_analysis = any([
        capa.get('fishBone'),
        capa.get('advancedRootCause'),
        capa.get('outcome')
    ])

    # Root Cause Section (based on priority)
    root_cause_section = ""
    if has_advanced_analysis:
        if capa.get('fishBone'):
            root_cause_section += f"\n**Fishbone Analysis**:\n{capa['fishBone']}\n"
        if capa.get('advancedRootCause'):
            root_cause_section += f"\n**Advanced Root Cause Analysis**:\n{capa['advancedRootCause']}\n"
        if capa.get('isIsNot'):
            root_cause_section += f"\n**Is / Is Not Analysis**:\n{capa['isIsNot']}\n"
        if capa.get('outcome'):
            root_cause_section += f"\n**Advanced Corrective Action (Outcome)**:\n{capa['outcome']}\n"
    else:
        root_cause_section += f"""
### **5 Why Analysis**:
1) **Why 1**: {capa.get('why1', 'N/A')}
2) **Why 2**: {capa.get('why2', 'N/A')}
3) **Why 3**: {capa.get('why3', 'N/A')}
4) **Why 4**: {capa.get('why4', 'N/A')}
5) **Why 5**: {capa.get('why5', 'N/A')}
"""

    # Build the complete prompt
    prompt = f"""
Convert the following CAPA (Corrective and Preventive Action) data into a meaningful and cohesive natural language report. Include location and entity information naturally within the text and maintain the order of key insights:
Ensure that the output is detailed and easy to understand, with proper flow and formatting, think of it like an author writing about the CAPA while maintaining all the information and its nuances.

If the total output exceeds **3000 characters**, ensure that every **3000-character segment** retains the **essence of the CAPA**, including:
- **CAPA Number**: {capa.get('serialNumber', 'N/A')}
- **Problem Statement**: {capa.get('title', 'N/A')}
- **Problem Description**: {capa.get('description', 'N/A')}
- **Location**: {capa.get('location', 'N/A')}
- **Entity**: {capa.get('entity', 'N/A')}

---

### **CAPA Details:**

**Organization**: {capa.get('organizationName', 'N/A')}
**Problem Statement**: {capa.get('title', 'N/A')}
**Problem Description**: {capa.get('description', 'N/A')}
**CAPA Origin**: {capa.get('origin', 'N/A')}
**CAPA Type**: {capa.get('type', 'N/A')}
**CAPA Number**: {capa.get('serialNumber', 'N/A')}
**Location**: {capa.get('location', 'N/A')}
**Entity**: {capa.get('entity', 'N/A')}
**CAPA Creator**: {capa.get('registeredBy', 'N/A')}
**CAPA Coordinator**: {capa.get('caraCoordinator', 'N/A')}
**CAPA Owner**: {capa.get('caraOwner', 'N/A')}
**Systems**: {capa.get('systems', 'N/A')}
**High Priority**: {capa.get('highPriority', 'N/A')}
**Defect Type**: {capa.get('defectType', 'N/A')}

**Containment Action**: {capa.get('containmentAction', 'N/A')}
**Planned Corrective Action**: {capa.get('correctiveAction', 'N/A')}
**Root Cause Analysis**: {capa.get('rootCauseAnalysis', 'N/A')}
**Actual Corrective Action**: {capa.get('actualCorrectiveAction', 'N/A')}

{root_cause_section}

### **Possible Causes**:
- Method: {capa.get('method', 'N/A')}
- Measurement: {capa.get('measurement', 'N/A')}
- Material: {capa.get('material', 'N/A')}
- Manpower: {capa.get('man', 'N/A')}
- Machine: {capa.get('machine', 'N/A')}
- Environment: {capa.get('environment', 'N/A')}

---

#### **Formatting Notes:**
If the total output exceeds **3000 characters**, insert a **reminder** every 3000 characters that repeats:  
- **CAPA Number**  
- **Problem Statement**  
- **Problem Description**  
- **Location**  
- **Entity**  

---

### **Example**:

For a CAPA in Bangalore, Karnataka, addressing a **A significant increase in the rejection rate of vials**, the report will start with:
*"In Bangalore, Karnataka, a significant increase in the rejection rate of vials was identified, increasing rejection rates from 1% to 6%. Defects primarily include cracks on the surface and incorrect dimensions, detected during the final quality inspection...."*

If the report is **longer than 3000 characters**, a reminder like:  
*"Continuing analysis for CAPA 24-BLRDept3-10: The issue in Bangalore, Karnataka involves rejection rate of vials due to surface cracks..."*  
will be inserted at intervals.

---

Ensure the report is **professional, clear, and semantically searchable**, with all insights naturally integrated and clearly structured.
"""

    
    openai_client = org_opeani(api_key=os.getenv("OPENAI_API_KEY"))
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert technical writer."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=4091,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        naturalText = response.choices[0].message.content
        return naturalText
    except Exception as e:
        print("Error generating suggestions:", e)
        return None

# def find_similar_capas(title, description, existing_capas, top_k=3):
#     query_text = f"{title}. {description}"
#     query_embedding = get_embedding(query_text)

#     if not query_embedding:
#         return []

#     similarities = []
#     for capa in existing_capas:
#         capa_text = f"{capa['title']}. {capa['description']}"
#         capa_embedding = get_embedding(capa_text)
#         if capa_embedding:
#             similarity = cosine_similarity([query_embedding], [capa_embedding])[0][0]
#             similarities.append((capa, similarity))

#     # Sort by similarity score (descending)
#     sorted_capas = sorted(similarities, key=lambda x: x[1], reverse=True)
#     similar_capas = [capa for capa, score in sorted_capas[:top_k]]

#     return similar_capas